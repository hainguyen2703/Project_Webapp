from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "Bilingual_Developer_Guide.docx"


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style_names = [
        "Normal",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Bullet",
        "List Number",
    ]
    size_map = {
        "Normal": 11,
        "Title": 20,
        "Subtitle": 12,
        "Heading 1": 16,
        "Heading 2": 14,
        "Heading 3": 12,
        "List Bullet": 11,
        "List Number": 11,
    }

    for style_name in style_names:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size_map[style_name])

    return doc


def add_title(doc: Document, title: str, subtitle: str) -> None:
    title_paragraph = doc.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True

    subtitle_paragraph = doc.add_paragraph(style="Subtitle")
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.add_run(subtitle)

    doc.add_paragraph("")


def add_bilingual_section(doc: Document, title: str, vn_lines: list[tuple[str, str]], en_lines: list[tuple[str, str]]) -> None:
    doc.add_heading(title, level=1)

    vn_label = doc.add_paragraph()
    vn_label.add_run("Tieng Viet").bold = True
    add_content_lines(doc, vn_lines)

    en_label = doc.add_paragraph()
    en_label.add_run("English").bold = True
    add_content_lines(doc, en_lines)

    doc.add_paragraph("")


def add_content_lines(doc: Document, lines: list[tuple[str, str]]) -> None:
    for kind, text in lines:
        if kind == "p":
            doc.add_paragraph(text)
        elif kind == "b":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "n":
            doc.add_paragraph(text, style="List Number")


def finalize_fonts(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Project_Webapp - Bilingual Developer Guide")


def build_sections() -> list[tuple[str, list[tuple[str, str]], list[tuple[str, str]]]]:
    return [
        (
            "1. Muc Dich Tai Lieu / Document Purpose",
            [
                (
                    "p",
                    "Tài liệu này được viết theo phong cách tài liệu kỹ thuật của ngôn ngữ lập trình, framework và thư viện. Mục tiêu là giúp người đọc không chỉ hiểu dự án đang có, mà còn biết cách viết thêm code, mở rộng tính năng và sử dụng công nghệ trong dự án một cách đúng hướng.",
                ),
                ("b", "Dùng như một tài liệu hướng dẫn lập trình cho dự án."),
                ("b", "Trình bày cả kiến trúc tổng quan và cách code từng phần."),
                (
                    "b",
                    "Hỗ trợ học cách dùng Python, Flask, Jinja, SQLite, APScheduler và pytest trong bối cảnh dự án này.",
                ),
            ],
            [
                (
                    "p",
                    "This document is written in the style of programming language, framework, and library documentation. Its goal is to help readers not only understand the current project, but also learn how to write new code, extend features, and use the project technologies in the correct way.",
                ),
                ("b", "Acts as a developer guide for the project."),
                ("b", "Explains both the overall architecture and the way each part is coded."),
                (
                    "b",
                    "Helps readers learn how Python, Flask, Jinja, SQLite, APScheduler, and pytest are used in this codebase.",
                ),
            ],
        ),
        (
            "2. Doi Tuong Doc / Target Audience",
            [
                ("b", "Sinh viên muốn đọc và hiểu dự án."),
                ("b", "Người phát triển muốn thêm route, service, template hoặc database logic."),
                (
                    "b",
                    "Người muốn học cách xây dựng một Flask web app có tính năng cá nhân hóa và notification.",
                ),
            ],
            [
                ("b", "Students who want to read and understand the project."),
                ("b", "Developers who want to add routes, services, templates, or database logic."),
                (
                    "b",
                    "Readers who want to learn how to build a Flask web app with personalization and notifications.",
                ),
            ],
        ),
        (
            "3. Cong Nghe Duoc Su Dung / Technologies Used",
            [
                ("b", "Python: ngôn ngữ chính của backend."),
                ("b", "Flask: framework web để định nghĩa route, request, response và render template."),
                ("b", "Flask-Login: quản lý đăng nhập và session của user."),
                ("b", "Jinja2: engine render HTML template."),
                (
                    "b",
                    "SQLite: database nhẹ lưu user, favourites, interests, snapshots và notifications.",
                ),
                ("b", "APScheduler: chạy tác vụ nền theo chu kỳ."),
                ("b", "python-arxiv: lấy bài báo từ arXiv."),
                ("b", "scikit-learn: tính similarity giữa các paper."),
                ("b", "pytest: kiểm thử logic và luồng chức năng."),
            ],
            [
                ("b", "Python: the main backend language."),
                ("b", "Flask: the web framework used for routes, requests, responses, and template rendering."),
                ("b", "Flask-Login: handles authentication and user sessions."),
                ("b", "Jinja2: the HTML template rendering engine."),
                (
                    "b",
                    "SQLite: a lightweight database for users, favourites, interests, snapshots, and notifications.",
                ),
                ("b", "APScheduler: runs background jobs on a schedule."),
                ("b", "python-arxiv: fetches papers from arXiv."),
                ("b", "scikit-learn: computes similarity between papers."),
                ("b", "pytest: tests logic and feature flows."),
            ],
        ),
        (
            "4. Cach Nhin Du An Nhu Tai Lieu Framework / How To Read The Project Like Framework Documentation",
            [
                (
                    "p",
                    "Nếu xem dự án này như tài liệu framework, cần nhớ rằng mỗi framework tốt đều trả lời bốn câu hỏi: nó là gì, nó được tổ chức thế nào, lập trình với nó ra sao, và muốn mở rộng thì mở rộng ở đâu. Dự án này cũng có thể đọc theo đúng bốn câu hỏi đó.",
                ),
                ("n", "Bắt đầu từ `src/app.py` để thấy điểm vào của hệ thống."),
                ("n", "Đọc `src/services/` để thấy business logic được tách ra như thế nào."),
                ("n", "Đọc `src/clients/` để thấy hệ thống lấy dữ liệu bên ngoài ra sao."),
                ("n", "Đọc `src/templates/` để thấy backend đưa dữ liệu ra giao diện như thế nào."),
                ("n", "Đọc `src/services/db.py` để hiểu dữ liệu được lưu ở đâu và theo cấu trúc nào."),
            ],
            [
                (
                    "p",
                    "If you read this project the way you would read framework documentation, remember that a good framework guide usually answers four questions: what it is, how it is organized, how to code with it, and where to extend it. This project can be understood using the same four questions.",
                ),
                ("n", "Start with `src/app.py` to see the system entry point."),
                ("n", "Read `src/services/` to understand how business logic is separated."),
                ("n", "Read `src/clients/` to see how external data is fetched."),
                ("n", "Read `src/templates/` to understand how backend data reaches the UI."),
                ("n", "Read `src/services/db.py` to understand where data is stored and how it is structured."),
            ],
        ),
        (
            "5. Python Trong Du An Nay / Python In This Project",
            [
                (
                    "p",
                    "Python được dùng làm ngôn ngữ backend chính. Code trong dự án theo hướng chia module rõ ràng, hàm nhỏ để tái sử dụng, và đặt business logic trong service thay vì viết hết trong route.",
                ),
                ("b", "Dùng type hint để code rõ ràng hơn."),
                ("b", "Dùng dataclass cho model dữ liệu đơn giản."),
                ("b", "Dùng function-level organization để mỗi logic dễ tìm và dễ test."),
                ("b", "Dùng import theo module để phân tách trách nhiệm rõ ràng."),
            ],
            [
                (
                    "p",
                    "Python is the main backend language in this project. The code follows a modular style, uses small reusable functions, and places business logic inside services rather than writing everything inside routes.",
                ),
                ("b", "Use type hints to make code clearer."),
                ("b", "Use dataclasses for simple data models."),
                ("b", "Use function-level organization so logic is easy to find and test."),
                ("b", "Use module-based imports to keep responsibilities separated."),
            ],
        ),
        (
            "6. Flask Framework Guide / Huong Dan Framework Flask",
            [
                (
                    "p",
                    "Flask là framework web trung tâm của dự án. Nó quản lý route, request, response, session và render template.",
                ),
                ("n", "Tạo route bằng `@app.route(...)`."),
                ("n", "Đọc tham số từ `request.args` hoặc `request.form`."),
                ("n", "Trả HTML bằng `render_template(...)`."),
                ("n", "Trả JSON bằng `jsonify(...)`."),
                ("n", "Chuyển hướng bằng `redirect(url_for(...))`."),
                (
                    "p",
                    "Nguyên tắc quan trọng trong dự án này là route chỉ nên điều phối, không nên chứa quá nhiều business logic.",
                ),
            ],
            [
                (
                    "p",
                    "Flask is the core web framework of the project. It manages routes, requests, responses, sessions, and template rendering.",
                ),
                ("n", "Create routes with `@app.route(...)`."),
                ("n", "Read parameters from `request.args` or `request.form`."),
                ("n", "Return HTML with `render_template(...)`."),
                ("n", "Return JSON with `jsonify(...)`."),
                ("n", "Redirect with `redirect(url_for(...))`."),
                (
                    "p",
                    "An important rule in this codebase is that routes should mainly coordinate work, not contain too much business logic.",
                ),
            ],
        ),
        (
            "7. Jinja Template Guide / Huong Dan Jinja Template",
            [
                (
                    "p",
                    "Jinja được dùng để render HTML ở server. Dữ liệu từ route được truyền vào template và hiển thị thông qua biến, vòng lặp, điều kiện và template inheritance.",
                ),
                ("b", "Dùng `base.html` làm layout gốc."),
                ("b", "Dùng `{% extends %}` để kế thừa layout."),
                ("b", "Dùng `{% block content %}` để chèn nội dung từng trang."),
                ("b", "Dùng `{{ ... }}` để in biến ra HTML."),
                ("b", "Dùng `{% if %}` và `{% for %}` để xử lý logic hiển thị."),
            ],
            [
                (
                    "p",
                    "Jinja is used to render HTML on the server. Data from routes is passed into templates and displayed through variables, loops, conditions, and template inheritance.",
                ),
                ("b", "Use `base.html` as the root layout."),
                ("b", "Use `{% extends %}` for layout inheritance."),
                ("b", "Use `{% block content %}` to inject page-specific content."),
                ("b", "Use `{{ ... }}` to print variables into HTML."),
                ("b", "Use `{% if %}` and `{% for %}` for display logic."),
            ],
        ),
        (
            "8. Database Guide With SQLite / Huong Dan Database Voi SQLite",
            [
                (
                    "p",
                    "SQLite được dùng để lưu trạng thái hệ thống. Trong dự án này, database không chỉ lưu user mà còn lưu favourites, interest selections, paper snapshots, related papers, notifications và metadata.",
                ),
                ("n", "Tất cả schema được tạo trong `init_db()` của `src/services/db.py`."),
                ("n", "Mọi thao tác đọc ghi nên đi qua helper function trong `db.py`."),
                ("n", "Không nên viết SQL trực tiếp trong template hoặc route."),
                ("n", "Nếu thêm tính năng mới, cần xem có cần thêm bảng, cột hay index mới hay không."),
            ],
            [
                (
                    "p",
                    "SQLite is used to store application state. In this project, the database stores not only users, but also favourites, interest selections, paper snapshots, related papers, notifications, and metadata.",
                ),
                ("n", "All schema creation is handled in `init_db()` inside `src/services/db.py`."),
                ("n", "All read and write operations should go through helper functions in `db.py`."),
                ("n", "Do not write SQL directly inside templates or routes."),
                ("n", "When adding a new feature, check whether you need a new table, column, or index."),
            ],
        ),
        (
            "9. Authentication And Session Guide / Huong Dan Xac Thuc Va Session",
            [
                (
                    "p",
                    "Hệ thống dùng Flask-Login để quản lý đăng nhập. Ngoài ra còn có session version và auth expiry để kiểm soát session chặt hơn.",
                ),
                ("b", "Đăng nhập: xác minh email và password hash."),
                ("b", "Throttle: giới hạn tạm thời sau nhiều lần đăng nhập sai."),
                ("b", "Logout: tăng session version để vô hiệu hóa session cũ."),
                ("b", "Before request: kiểm tra session và đồng bộ interests của user."),
            ],
            [
                (
                    "p",
                    "The system uses Flask-Login for authentication. It also uses a session version and auth expiry to control session validity more strictly.",
                ),
                ("b", "Login: verifies email and password hash."),
                ("b", "Throttle: temporarily limits repeated failed login attempts."),
                ("b", "Logout: bumps the session version to invalidate older sessions."),
                ("b", "Before request: checks session state and reconciles user interests."),
            ],
        ),
        (
            "10. Discovery Flow Guide / Huong Dan Luong Discovery",
            [
                (
                    "p",
                    "Discovery là tính năng cốt lõi của dự án. Mục tiêu của nó là lấy bài báo từ arXiv, lọc, xếp hạng, phân trang và hiển thị cho user.",
                ),
                ("n", "Route nhận query hoặc tạo query mặc định từ interests."),
                ("n", "`fetch_items()` gọi arXiv client."),
                ("n", "Kết quả được validate và rank theo relevance cộng recency."),
                ("n", "Các snapshot paper được lưu vào database."),
                ("n", "Related papers được tính để phục vụ trang detail."),
            ],
            [
                (
                    "p",
                    "Discovery is the core feature of the project. Its purpose is to fetch papers from arXiv, filter them, rank them, paginate them, and display them to the user.",
                ),
                ("n", "The route receives a query or builds a default query from user interests."),
                ("n", "`fetch_items()` calls the arXiv client."),
                ("n", "Results are validated and ranked by relevance plus recency."),
                ("n", "Paper snapshots are stored in the database."),
                ("n", "Related papers are computed for the detail page."),
            ],
        ),
        (
            "11. How To Add A New Feature / Cach Them Tinh Nang Moi",
            [
                (
                    "p",
                    "Khi thêm một tính năng mới, nên đi theo luồng có cấu trúc thay vì sửa ngẫu nhiên ở nhiều nơi.",
                ),
                ("n", "Xác định route mới hay mở rộng route cũ trong `src/app.py`."),
                ("n", "Tách logic xử lý ra service nếu logic không quá nhỏ."),
                ("n", "Nếu cần lưu dữ liệu, thêm helper function trong `src/services/db.py`."),
                ("n", "Nếu cần hiển thị giao diện, thêm hoặc sửa template trong `src/templates/`."),
                ("n", "Nếu tính năng có nguy cơ hồi quy, thêm test phù hợp."),
            ],
            [
                (
                    "p",
                    "When adding a new feature, follow a structured flow instead of making scattered edits across many files.",
                ),
                ("n", "Decide whether to add a new route or extend an existing one in `src/app.py`."),
                ("n", "Move logic into a service if the logic is more than trivial."),
                ("n", "If persistence is needed, add helper functions in `src/services/db.py`."),
                ("n", "If UI is needed, add or update a template in `src/templates/`."),
                ("n", "If the feature risks regression, add an appropriate test."),
            ],
        ),
        (
            "12. Framework Style Coding Rules / Quy Tac Viet Code Theo Kieu Framework",
            [
                ("b", "Route mỏng, service dày: route nên ngắn, service nên chứa logic."),
                ("b", "Model rõ ràng: dữ liệu nên có hình dạng ổn định và dễ validate."),
                ("b", "Template tập trung vào hiển thị, không giải quyết business logic lớn."),
                ("b", "Database helper là điểm truy cập dữ liệu chính."),
                ("b", "Tên hàm, tên biến và tên module nên nói rõ mục đích."),
            ],
            [
                ("b", "Thin routes, rich services: routes should stay short while services hold logic."),
                ("b", "Clear models: data should have a stable, validated shape."),
                ("b", "Templates should focus on presentation, not large business rules."),
                ("b", "Database helpers should be the main access point for persistence."),
                ("b", "Function names, variable names, and module names should clearly express intent."),
            ],
        ),
        (
            "13. Testing Guide / Huong Dan Kiem Thu",
            [
                (
                    "p",
                    "Dự án có unit test và integration test. Đây cũng là cách mà tài liệu framework thường dạy: test từng đơn vị logic, sau đó test luồng người dùng.",
                ),
                ("b", "Unit test kiểm tra hàm, model và service nhỏ."),
                ("b", "Integration test kiểm tra route và luồng tính năng từ đầu đến cuối."),
                ("b", "Khi sửa auth, discovery, registration hoặc database logic, nên xem test liên quan trước."),
            ],
            [
                (
                    "p",
                    "The project contains both unit tests and integration tests. This follows the same pattern often recommended by framework documentation: test the small logic units first, then test full user flows.",
                ),
                ("b", "Unit tests check functions, models, and focused services."),
                ("b", "Integration tests check routes and feature flows end to end."),
                ("b", "When changing auth, discovery, registration, or database logic, review the related tests first."),
            ],
        ),
        (
            "14. Suggested Learning Path / Lo Trinh Hoc De Nhanh Hieu Du An",
            [
                ("n", "Đọc `src/app.py` để thấy toàn bộ map route và luồng request."),
                ("n", "Đọc `src/services/discovery_service.py` để hiểu tính năng cốt lõi."),
                ("n", "Đọc `src/clients/arxiv_client.py` để hiểu nguồn dữ liệu."),
                ("n", "Đọc `src/services/db.py` để thấy hệ thống lưu gì và lưu như thế nào."),
                ("n", "Đọc template để thấy giao diện dùng dữ liệu ra sao."),
                ("n", "Cuối cùng đọc tests để thấy hệ thống được kỳ vọng hoạt động như thế nào."),
            ],
            [
                ("n", "Read `src/app.py` to understand the route map and request flow."),
                ("n", "Read `src/services/discovery_service.py` to understand the core feature."),
                ("n", "Read `src/clients/arxiv_client.py` to understand the data source."),
                ("n", "Read `src/services/db.py` to understand what is stored and how it is stored."),
                ("n", "Read the templates to see how the UI consumes data."),
                ("n", "Finally, read the tests to understand expected behavior."),
            ],
        ),
        (
            "15. Final Summary / Tong Ket Cuoi",
            [
                (
                    "p",
                    "Nếu xem dự án này như một bộ tài liệu lập trình, thì nó dạy cho người đọc ba điều cùng lúc: cách đọc một Flask monolith, cách tổ chức business logic thành service, và cách ghép source dữ liệu bên ngoài với giao diện web và database. Vì vậy, tài liệu này có thể được dùng như hướng dẫn đọc code, hướng dẫn mở rộng tính năng, và hướng dẫn học công nghệ đã được sử dụng trong dự án.",
                ),
            ],
            [
                (
                    "p",
                    "If you treat this project like a programming guide, it teaches three things at once: how to read a Flask monolith, how to organize business logic into services, and how to connect an external data source with a web interface and a database. For that reason, this document can be used as a code-reading guide, a feature-extension guide, and a technology usage guide for the stack used in the project.",
                ),
            ],
        ),
    ]


def main() -> None:
    doc = configure_document()
    add_title(
        doc,
        "Hướng Dẫn Phát Triển Song Ngữ / Bilingual Developer Guide",
        "Vietnamese and English are shown together section by section for project study and implementation.",
    )

    for title, vn_lines, en_lines in build_sections():
        add_bilingual_section(doc, title, vn_lines, en_lines)

    finalize_fonts(doc)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
